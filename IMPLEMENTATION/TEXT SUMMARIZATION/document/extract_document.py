class Document:
    
    def extractPDFDocument(self, inputFileName):
        
        import PyPDF2 

        input_text = ""
        
        pdfFileobj = open(inputFileName,'rb')

        pdfReader = PyPDF2.PdfFileReader(pdfFileobj)
        
        for i in range(pdfReader.numPages):
            pageobj = pdfReader.getPage(i)

            with open('text.txt','a')as file:
                content  = pageobj.extractText()
                file.write(content)
                        
            input_text += pageobj.extractText()            

        pdfFileobj.close()

        return input_text

    def extractWordDocument(self, inputFileName):
        import docx

        input_text = ""

        doc = docx.Document(inputFileName)
        all_paras = doc.paragraphs

        for para in all_paras:
            input_text += para.text

        return input_text        
